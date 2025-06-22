from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import datetime
from dataclasses import dataclass
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, PageBreak
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import textwrap
import html
import re


@dataclass
class OutputFile:
    file_path: str
    content: str


class OutputFormatter(ABC):
    def __init__(self, config_manager):
        self.config_manager = config_manager
    
    @abstractmethod
    def format_output(self, files: List[OutputFile], output_path: Path, prompt: str = "") -> None:
        pass
    
    def get_separator(self) -> str:
        return self.config_manager.get('output_settings.file_separator', '=== FILE: {filepath} ===')
    
    def get_prompt_placeholder(self) -> str:
        return self.config_manager.get('output_settings.prompt_placeholder', '[PROMPT]')
    
    def get_content_placeholder(self) -> str:
        return self.config_manager.get('output_settings.content_placeholder', '[CONTENT]')


class TextOutputFormatter(OutputFormatter):
    def format_output(self, files: List[OutputFile], output_path: Path, prompt: str = "") -> None:
        encoding = self.config_manager.get('encoding', 'utf-8')
        
        with open(output_path, 'w', encoding=encoding) as f:
            # Write header
            f.write(f"# CodeFuser Output\n")
            f.write(f"# Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# Total Files: {len(files)}\n")
            f.write("\n" + "="*80 + "\n\n")
            
            # Write prompt if provided (once at the beginning)
            if prompt:
                f.write(f"{self.get_prompt_placeholder()}\n")
                f.write(f"{prompt}\n")
                f.write("\n" + "-"*80 + "\n\n")
            
            for idx, file_data in enumerate(files):
                # Write file separator
                separator = self.get_separator().format(filepath=file_data.file_path)
                f.write(f"\n{separator}\n")
                
                # Write content
                f.write(f"\n{self.get_content_placeholder()}\n")
                f.write(file_data.content)
                
                # Add spacing between files
                if idx < len(files) - 1:
                    f.write("\n\n")


class DocxOutputFormatter(OutputFormatter):
    def format_output(self, files: List[OutputFile], output_path: Path, prompt: str = "") -> None:
        doc = Document()
        
        # Add title
        title = doc.add_heading('CodeFuser Output', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Files: {len(files)}")
        doc.add_paragraph("")
        
        # Add prompt if provided (once at the beginning)
        if prompt:
            prompt_heading = doc.add_heading(self.get_prompt_placeholder(), level=2)
            prompt_para = doc.add_paragraph(prompt)
            prompt_para.style.font.size = Pt(11)
            prompt_para.style.font.color.rgb = RGBColor(0, 0, 139)
            doc.add_paragraph("-" * 80)
            doc.add_paragraph("")
        
        for idx, file_data in enumerate(files):
            # Add file header
            file_header = doc.add_heading(file_data.file_path, level=2)
            
            # Add content
            content_heading = doc.add_heading(self.get_content_placeholder(), level=3)
            
            # Split content into smaller chunks for better handling
            content_lines = file_data.content.split('\n')
            for line in content_lines:
                if line.strip():
                    para = doc.add_paragraph(line)
                    para.style.font.name = 'Courier New'
                    para.style.font.size = Pt(9)
                else:
                    doc.add_paragraph("")
            
            # Add page break between files (except for the last one)
            if idx < len(files) - 1:
                doc.add_page_break()
        
        # Save document
        doc.save(str(output_path))


class PdfOutputFormatter(OutputFormatter):
    def __init__(self, config_manager):
        super().__init__(config_manager)
        self._register_fonts()
    
    def _register_fonts(self):
        # Try to register monospace fonts
        try:
            # This is a placeholder - in production, you'd want to handle font paths properly
            pass
        except:
            pass
    
    def format_output(self, files: List[OutputFile], output_path: Path, prompt: str = "") -> None:
        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='darkblue',
            alignment=TA_LEFT,
        )
        
        file_header_style = ParagraphStyle(
            'FileHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor='darkgreen',
            spaceAfter=12,
        )
        
        code_style = ParagraphStyle(
            'Code',
            parent=styles['Code'],
            fontSize=8,
            fontName='Courier',
            leftIndent=0,
            rightIndent=0,
            spaceAfter=6,
        )
        
        # Add title
        elements.append(Paragraph("CodeFuser Output", title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Add metadata
        metadata = f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>Total Files: {len(files)}"
        elements.append(Paragraph(metadata, styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
        
        # Add prompt if provided (once at the beginning)
        if prompt:
            elements.append(Paragraph(f"<b>{self.get_prompt_placeholder()}</b>", styles['Heading2']))
            elements.append(Paragraph(prompt, styles['Italic']))
            elements.append(Spacer(1, 0.3*inch))
            elements.append(Paragraph("-" * 80, styles['Normal']))
            elements.append(Spacer(1, 0.3*inch))
        
        # Process files
        for idx, file_data in enumerate(files):
            # Add file header
            elements.append(Paragraph(file_data.file_path, file_header_style))
            
            # Add content header
            elements.append(Paragraph(f"<b>{self.get_content_placeholder()}</b>", styles['Normal']))
            elements.append(Spacer(1, 0.05*inch))
            
            # Add content - wrap long lines
            wrapped_content = self._wrap_content(file_data.content, 100)
            content_pre = Preformatted(wrapped_content, code_style)
            elements.append(content_pre)
            
            # Add page break between files (except for the last one)
            if idx < len(files) - 1:
                elements.append(PageBreak())
        
        # Build PDF
        doc.build(elements)
    
    def _wrap_content(self, content: str, width: int = 100) -> str:
        lines = content.split('\n')
        wrapped_lines = []
        
        for line in lines:
            if len(line) > width:
                wrapped = textwrap.wrap(line, width=width, break_long_words=True, break_on_hyphens=False)
                wrapped_lines.extend(wrapped)
            else:
                wrapped_lines.append(line)
        
        return '\n'.join(wrapped_lines)


class HtmlOutputFormatter(OutputFormatter):
    def format_output(self, files: List[OutputFile], output_path: Path, prompt: str = "") -> None:
        html_content = self._generate_html(files, prompt)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def _generate_html(self, files: List[OutputFile], prompt: str = "") -> str:
        html_template = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CodeFuser Output</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/themes/prism-tomorrow.min.css">
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/components/prism-core.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/prism/1.29.0/plugins/autoloader/prism-autoloader.min.js"></script>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', sans-serif;
            line-height: 1.6;
            color: #333;
            background: #f8f9fa;
            padding: 20px;
        }}
        
        .container {{
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 12px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.1);
            overflow: hidden;
        }}
        
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }}
        
        .header h1 {{
            font-size: 2.5rem;
            margin-bottom: 10px;
            font-weight: 300;
        }}
        
        .header .meta {{
            opacity: 0.9;
            font-size: 1.1rem;
        }}
        
        .prompt-section {{
            background: #e3f2fd;
            padding: 25px;
            margin: 20px;
            border-radius: 8px;
            border-left: 4px solid #2196F3;
        }}
        
        .prompt-section h2 {{
            color: #1976D2;
            margin-bottom: 15px;
            font-size: 1.4rem;
        }}
        
        .prompt-content {{
            background: white;
            padding: 20px;
            border-radius: 6px;
            white-space: pre-wrap;
            font-family: 'SF Mono', Monaco, 'Cascadia Code', monospace;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }}
        
        .file-section {{
            margin: 30px 20px;
            border: 1px solid #e1e5e9;
            border-radius: 8px;
            overflow: hidden;
        }}
        
        .file-header {{
            background: #f8f9fa;
            padding: 15px 20px;
            border-bottom: 1px solid #e1e5e9;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }}
        
        .file-path {{
            font-family: 'SF Mono', Monaco, 'Cascadia Code', monospace;
            font-weight: 600;
            color: #0366d6;
            font-size: 1.1rem;
        }}
        
        .file-info {{
            display: flex;
            gap: 15px;
            font-size: 0.9rem;
            color: #586069;
        }}
        
        .file-content {{
            position: relative;
        }}
        
        .copy-button {{
            position: absolute;
            top: 10px;
            right: 15px;
            background: #24292e;
            color: white;
            border: none;
            padding: 8px 12px;
            border-radius: 6px;
            font-size: 0.8rem;
            cursor: pointer;
            opacity: 0.8;
            transition: opacity 0.2s;
        }}
        
        .copy-button:hover {{
            opacity: 1;
        }}
        
        .copy-button:active {{
            background: #0366d6;
        }}
        
        pre[class*="language-"] {{
            margin: 0 !important;
            border-radius: 0 !important;
            font-size: 0.9rem;
            line-height: 1.5;
        }}
        
        code[class*="language-"] {{
            font-family: 'SF Mono', Monaco, 'Cascadia Code', 'Roboto Mono', monospace !important;
        }}
        
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            padding: 20px;
            background: #f8f9fa;
            border-top: 1px solid #e1e5e9;
        }}
        
        .stat-item {{
            text-align: center;
            padding: 15px;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }}
        
        .stat-value {{
            font-size: 2rem;
            font-weight: 600;
            color: #0366d6;
        }}
        
        .stat-label {{
            color: #586069;
            font-size: 0.9rem;
            margin-top: 5px;
        }}
        
        .toc {{
            background: #f6f8fa;
            padding: 20px;
            margin: 20px;
            border-radius: 8px;
            border: 1px solid #e1e5e9;
        }}
        
        .toc h3 {{
            margin-bottom: 15px;
            color: #24292e;
        }}
        
        .toc ul {{
            list-style: none;
            padding-left: 0;
        }}
        
        .toc li {{
            padding: 5px 0;
        }}
        
        .toc a {{
            color: #0366d6;
            text-decoration: none;
            font-family: 'SF Mono', Monaco, monospace;
            font-size: 0.9rem;
        }}
        
        .toc a:hover {{
            text-decoration: underline;
        }}
        
        @media (max-width: 768px) {{
            body {{
                padding: 10px;
            }}
            
            .header h1 {{
                font-size: 2rem;
            }}
            
            .stats {{
                grid-template-columns: repeat(2, 1fr);
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            <h1>🚀 CodeFuser Output</h1>
            <div class="meta">
                Generated on {timestamp}<br>
                {file_count} files processed
            </div>
        </div>
        
        <!-- Statistics -->
        <div class="stats">
            <div class="stat-item">
                <div class="stat-value">{file_count}</div>
                <div class="stat-label">Total Files</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">{total_size}</div>
                <div class="stat-label">Total Size</div>
            </div>
            <div class="stat-item">
                <div class="stat-value">{language_count}</div>
                <div class="stat-label">Languages</div>
            </div>
        </div>
        
        {prompt_section}
        
        <!-- Table of Contents -->
        <div class="toc">
            <h3>📋 Table of Contents</h3>
            <ul>
                {toc_items}
            </ul>
        </div>
        
        <!-- File Contents -->
        {file_contents}
    </div>
    
    <script>
        // Copy to clipboard functionality
        function copyToClipboard(text) {{
            navigator.clipboard.writeText(text).then(() => {{
                // Visual feedback
                event.target.textContent = '✓ Copied!';
                setTimeout(() => {{
                    event.target.textContent = '📋 Copy';
                }}, 2000);
            }});
        }}
        
        // Add copy buttons to all code blocks
        document.addEventListener('DOMContentLoaded', function() {{
            document.querySelectorAll('pre[class*="language-"]').forEach(function(block) {{
                const button = document.createElement('button');
                button.className = 'copy-button';
                button.textContent = '📋 Copy';
                button.onclick = () => copyToClipboard(block.textContent);
                
                block.parentElement.style.position = 'relative';
                block.parentElement.appendChild(button);
            }});
        }});
    </script>
</body>
</html>"""
        
        # Generate content parts
        timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Statistics
        total_size = sum(len(f.content.encode('utf-8')) for f in files)
        languages = set()
        for f in files:
            ext = Path(f.file_path).suffix.lower()
            lang = self._get_language_from_extension(ext)
            if lang:
                languages.add(lang)
        
        size_str = self._format_size(total_size)
        
        # Prompt section
        prompt_section = ""
        if prompt:
            prompt_section = f"""
        <div class="prompt-section">
            <h2>🎯 Prompt Instructions</h2>
            <div class="prompt-content">{html.escape(prompt)}</div>
        </div>"""
        
        # Table of contents
        toc_items = []
        for f in files:
            file_id = self._sanitize_id(f.file_path)
            toc_items.append(f'<li><a href="#{file_id}">{html.escape(f.file_path)}</a></li>')
        
        # File contents
        file_contents_html = []
        for f in files:
            file_id = self._sanitize_id(f.file_path)
            language = self._get_language_from_extension(Path(f.file_path).suffix.lower())
            
            file_size = len(f.content.encode('utf-8'))
            line_count = len(f.content.split('\n'))
            
            file_html = f"""
        <div class="file-section" id="{file_id}">
            <div class="file-header">
                <div class="file-path">📄 {html.escape(f.file_path)}</div>
                <div class="file-info">
                    <span>{self._format_size(file_size)}</span>
                    <span>{line_count} lines</span>
                    <span>{language or 'text'}</span>
                </div>
            </div>
            <div class="file-content">
                <pre><code class="language-{language or 'text'}">{html.escape(f.content)}</code></pre>
            </div>
        </div>"""
            file_contents_html.append(file_html)
        
        # Fill template
        return html_template.format(
            timestamp=timestamp,
            file_count=len(files),
            total_size=size_str,
            language_count=len(languages),
            prompt_section=prompt_section,
            toc_items='\n                '.join(toc_items),
            file_contents='\n        '.join(file_contents_html)
        )
    
    def _get_language_from_extension(self, ext: str) -> str:
        """Map file extension to Prism.js language identifier"""
        lang_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.jsx': 'jsx',
            '.tsx': 'tsx',
            '.html': 'html',
            '.htm': 'html',
            '.css': 'css',
            '.scss': 'scss',
            '.sass': 'sass',
            '.json': 'json',
            '.xml': 'xml',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.md': 'markdown',
            '.java': 'java',
            '.cs': 'csharp',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.hpp': 'cpp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bash': 'bash',
            '.ps1': 'powershell',
            '.r': 'r',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala'
        }
        return lang_map.get(ext, 'text')
    
    def _format_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
    
    def _sanitize_id(self, filepath: str) -> str:
        """Create a valid HTML ID from filepath"""
        return re.sub(r'[^a-zA-Z0-9_-]', '_', filepath)


class OutputManager:
    def __init__(self, config_manager):
        self.config_manager = config_manager
        self.formatters = {
            'txt': TextOutputFormatter(config_manager),
            'docx': DocxOutputFormatter(config_manager),
            'pdf': PdfOutputFormatter(config_manager),
            'html': HtmlOutputFormatter(config_manager)
        }
    
    def create_output(
        self,
        files: List[Dict[str, Any]],
        output_path: Path,
        format: str,
        prompt: str = ""
    ) -> Path:
        
        if format not in self.formatters:
            raise ValueError(f"Unsupported output format: {format}")
        
        # Convert file data to OutputFile objects
        output_files = []
        encoding = self.config_manager.get('encoding', 'utf-8')
        
        for file_info in files:
            try:
                with open(file_info['path'], 'r', encoding=encoding) as f:
                    content = f.read()
                
                output_file = OutputFile(
                    file_path=file_info['relative_path'],
                    content=content
                )
                output_files.append(output_file)
            except Exception as e:
                # Log error but continue with other files
                print(f"Error reading file {file_info['path']}: {e}")
                continue
        
        # Ensure output path has correct extension
        output_path = output_path.with_suffix(f'.{format}')
        
        # Format and save output
        formatter = self.formatters[format]
        formatter.format_output(output_files, output_path, prompt)
        
        return output_path
    
    def get_available_formats(self) -> List[str]:
        return list(self.formatters.keys())
    
    def estimate_output_size(self, files: List[Dict[str, Any]], format: str) -> int:
        # Rough estimation based on file count and format
        total_size = sum(f.get('size', 0) for f in files)
        
        # Add overhead based on format
        overhead_multipliers = {
            'txt': 1.1,   # 10% overhead for separators and prompts
            'docx': 1.5,  # 50% overhead for XML structure
            'pdf': 2.0,   # 100% overhead for PDF structure
            'html': 3.0   # 200% overhead for HTML structure and styling
        }
        
        multiplier = overhead_multipliers.get(format, 1.2)
        return int(total_size * multiplier)